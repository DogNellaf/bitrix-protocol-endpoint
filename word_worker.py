from docx import Document
import copy
import random
from utils import parse_range

def clone_row(table, row_idx):
    """
    Клонирует строку таблицы по индексу row_idx (например, первую строку, чтобы сохранить стиль).
    Добавляет клонированную строку в конец таблицы и возвращает её.
    """
    new_row_xml = copy.deepcopy(table.rows[row_idx]._tr)
    table._tbl.append(new_row_xml)
    return table.rows[-1]

def clear_row_text(row):
    """
    Очищает текст во всех ячейках строки, но сохраняет структуру и форматирование.
    """
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            # В каждом параграфе обнуляем текст у всех run
            for run in paragraph.runs:
                run.text = ""

def update_cell_text(cell, text, template_cell):
    """
    Обновляет текст в ячейке cell, используя форматирование из соответствующей template_cell.
    Если в ячейке уже есть хотя бы один run, обновляет его текст и копирует настройки шрифта,
    за исключением жирности, которая отключается.
    Если нет – создаёт run с таким же форматированием.
    """
    if cell.paragraphs:
        para = cell.paragraphs[0]
    else:
        para = cell.add_paragraph()

    templ_run = None
    if template_cell.paragraphs and template_cell.paragraphs[0].runs:
        templ_run = template_cell.paragraphs[0].runs[0]

    if para.runs:
        target_run = para.runs[0]
        target_run.text = text
        if templ_run:
            target_run.font.size = templ_run.font.size
            target_run.font.name = templ_run.font.name
            target_run.font.bold = False  # Отключаем жирное начертание
            target_run.font.italic = templ_run.font.italic
            target_run.font.underline = templ_run.font.underline
    else:
        new_run = para.add_run(text)
        if templ_run:
            new_run.font.size = templ_run.font.size
            new_run.font.name = templ_run.font.name
            new_run.font.bold = False  # Отключаем жирное начертание
            new_run.font.italic = templ_run.font.italic
            new_run.font.underline = templ_run.font.underline

def add_row_with_template(table):
    """
    Добавляет новую строку в таблицу, клонируя первую строку таблицы (row 0).
    Сохраняется форматирование первой строки.
    """
    template_row = table.rows[0]
    new_row = clone_row(table, 0)
    clear_row_text(new_row)
    return new_row, template_row

def replace_text_in_paragraph(paragraph, placeholder, replacement):
    """Заменяет placeholder в каждом run абзаца, сохраняя стиль."""
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)

def change_placeholder_in_table(table, placeholders, values):
    """Проходит по всем ячейкам таблицы и заменяет placeholder в каждом параграфе."""
    for row in table.rows:
        for cell in row.cells:
            for placeholder in placeholders:
                placeholder = placeholder
                if placeholder in cell.text:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, placeholder, values[placeholder])

def fill_template(template_path, data, bitrix_fields):
    """
    Открывает шаблон Word и заменяет в нём плейсхолдеры на данные из google_data и bitrix_fields.
    """
    # Открываем шаблон
    doc = Document(template_path)

    # Заполняем данные под печатью (первая таблица)
    top_table = doc.tables[0]
    placeholders = ['{UfCrm51741286877382}', '{UfCrm51741286861298}']
    change_placeholder_in_table(top_table, placeholders, bitrix_fields)

    # Заполняем поле "ПРОТОКОЛ ИСПЫТАНИЙ" (вторая таблица)
    title_table = doc.tables[1]
    placeholders = [
        '{UfCrm51741285071989}',
        '{UfCrm51741286877382}',
        '{UfCrm51741285600655}',
        '{UfCrm51741799909007}',
        '{UfCrm51741285642541}',
        '{UfCrm51741787685Title}',
        '{UfCrm51741787685UfCrm1741797658252}',
        '{UfCrm51741787685UfCrm1741797676159}',
        '{UfCrm51741798185411}',
        '{UfCrm51741798069674}',
        '{UfCrm51741798090701}',
        '{UfCrm51741285472446}',
        '{UfCrm51741285487155}'
    ]
    change_placeholder_in_table(title_table, placeholders, bitrix_fields)

    # Заполняем параметры в начале второй страницы (третья таблица)
    second_page_table = doc.tables[2]
    placeholders = [
        '{UfCrm51741285539701}',
        '{UfCrm51741285553037}',
        '{UfCrm51741285167890}',
        '{UfCrm51741286822180}'
    ]
    change_placeholder_in_table(second_page_table, placeholders, bitrix_fields)
    
    # Заполняем "Нормативно-техническая документация на методы испытаний:"
    methods_text = "\n".join(data["Нормативно-техническая документация"])
    placeholder = "{methods}"
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            replace_text_in_paragraph(paragraph, placeholder, methods_text)
    
    # Заполняем таблицу "Используемое оборудование"
    equipments_table = doc.tables[-2]
    index = 1
    for equipment in data["Используемое оборудование"]:
        new_row, template_row = add_row_with_template(equipments_table)
        # Обновляем текст ячеек с копированием форматирования из соответствующих ячеек шаблона
        update_cell_text(new_row.cells[0], str(index), template_row.cells[0])
        update_cell_text(new_row.cells[1], equipment['Наименование, заводской номер, инвентаризационный номер'], template_row.cells[1])
        update_cell_text(new_row.cells[2], equipment['Дата окончания поверки, аттестации'], template_row.cells[2])
        index += 1
        
    # Заполняем таблицу "Результаты испытаний"
    results_table = doc.tables[-2]
    for result in data["Результаты испытаний"]:
        new_row, template_row = add_row_with_template(results_table)
        update_cell_text(new_row.cells[0], result['Показатель'], template_row.cells[0])
        update_cell_text(new_row.cells[1], result['Единицы измерений '], template_row.cells[1])
        update_cell_text(new_row.cells[2], result['Методы испытаний '], template_row.cells[2])
        update_cell_text(new_row.cells[3], result['Норма по НД'], template_row.cells[3])

        random_value = f"{parse_range(result['Результат испытаний '])}"
        update_cell_text(new_row.cells[4], random_value, template_row.cells[4])
    
    # Заполняем нижний колонтитул
    footer = doc.sections[0].footer.paragraphs[0]
    placeholders = [
        '{UfCrm51741286877382}',
        '{UfCrm51741286861298}'
    ]
    for placeholder in placeholders:
        if placeholder in footer.text:
            replace_text_in_paragraph(footer, placeholder, bitrix_fields[placeholder])

    return doc

def save_word_document(doc, filename):
    doc.save(filename)
    return filename
